"""
Test script to verify our DOCX file extraction from resumeBuilder.py
"""
import os
import sys
import docx
import traceback

def extract_text_from_docx(docx_path):
    """Extract text from a DOCX file"""
    try:
        print(f"Extracting text from DOCX: {docx_path}")
        doc = docx.Document(docx_path)
        
        # Extract text from paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        
        content = "\n".join(paragraphs)
        print(f"Successfully extracted {len(content)} characters from DOCX")
        return content
    except Exception as e:
        print(f"Error processing DOCX: {str(e)}")
        traceback.print_exc()
        return None

def main():
    docx_file = "test_template.docx"
    
    if not os.path.exists(docx_file):
        print(f"Error: {docx_file} does not exist")
        return 1
    
    content = extract_text_from_docx(docx_file)
    
    if content:
        print("\nExtracted content preview:")
        print("-" * 40)
        preview = "\n".join(content.split("\n")[:10])
        print(preview)
        print("-" * 40)
        print("Successfully read DOCX file!")
        return 0
    else:
        print("Failed to extract content from DOCX file")
        return 1

if __name__ == "__main__":
    sys.exit(main()) 