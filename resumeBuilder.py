from groq import Groq, GroqError
import datetime
from http.server import HTTPServer, BaseHTTPRequestHandler
import json
import os
import cgi
from urllib.parse import parse_qs
import traceback
import tempfile
import re
import time
import io
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import base64
import logging
import shutil
from pathlib import Path
import docx
import PyPDF2
from PIL import Image

# Configure logging
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('resume_server.log')
    ]
)
logger = logging.getLogger('ResumeServer')

def sanitize_text(text):
    if not isinstance(text, str):
        return ""
    
    # Remove markdown formatting
    text = re.sub(r'\*+', ' ', text)
    text = re.sub(r'#+\s*', '', text)
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    text = re.sub(r'!\[.*?\]\(.*?\)', '', text)
    text = re.sub(r'---+', '', text)
    text = re.sub(r'`{3}.*?`{3}', '', text, flags=re.DOTALL)
    text = re.sub(r'>{1,}\s*', '', text)
    text = re.sub(r'[-*+]\s', '', text)
    text = re.sub(r'\d+\.\s', '', text)
    
    # Remove special characters
    text = re.sub(r'[`~!@#$%^&*()_+={}\[\]|\\:;"\'<>,.?/]', ' ', text)
    
    # Normalize whitespace
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = text.strip()
    
    return text

def clean_text_for_json(text):
    if not isinstance(text, str):
        return ""
        
    # Remove common prefixes
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        if not re.match(r'^(Here\'s|Here is|Here are|Results|Analysis|Recommendations|Note:|Please find)', line.strip(), re.IGNORECASE):
            cleaned_lines.append(line)
    text = '\n'.join(cleaned_lines)
    
    # Remove formatting
    text = re.sub(r'\*\*.*?\*\*', '', text)
    text = re.sub(r'\*.*?\*', '', text)
    text = re.sub(r'`.*?`', '', text)
    text = re.sub(r'#+\s*', '', text)
    text = re.sub(r'\[.*?\]\(.*?\)', '', text)
    text = re.sub(r'!\[.*?\]\(.*?\)', '', text)
    text = re.sub(r'[-*+]\s+', '', text)
    text = re.sub(r'\d+\.\s+', '', text)
    
    # Clean up
    text = re.sub(r'[^a-zA-Z0-9\s\n]', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
    text = text.strip()
    
    return text

def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file"""
    try:
        logger.info(f"Extracting text from PDF: {pdf_path}")
        text = ""
        with open(pdf_path, 'rb') as file:
            try:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
            except Exception as e:
                logger.error(f"Error extracting PDF text: {str(e)}")
                return "Error extracting PDF content"
            
        return text
    except Exception as e:
        logger.error(f"Error opening PDF: {str(e)}")
        return "Error opening PDF file"

def process_template(template_path):
    """Process template image and return its dimensions and format for context information"""
    try:
        logger.info(f"Processing template image: {template_path}")
        with Image.open(template_path) as img:
            # Get image dimensions and format
            width, height = img.size
            format = img.format.lower() if img.format else 'unknown'
            
            logger.info(f"Processed template image: {width}x{height} {format}")
            
            return {
                'width': width,
                'height': height,
                'format': format,
                'mode': img.mode
            }
    except Exception as e:
        logger.error(f"Error processing template image: {str(e)}")
        return None

def safe_read_file(file_path):
    """Safely read file content regardless of type and encoding"""
    logger.info(f"Reading file: {file_path}")
    
    # Check if it's a PDF file
    if file_path.lower().endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    
    # Check if it's an image file
    try:
        with Image.open(file_path) as img:
            logger.info(f"File is an image: {file_path}")
            return "Image content detected. No text to extract."
    except:
        logger.debug(f"File is not an image: {file_path}")
    
    # Try different encodings for text files
    encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                content = file.read()
                logger.info(f"Successfully read file with {encoding} encoding")
                return content
        except UnicodeDecodeError:
            logger.debug(f"Failed to decode with {encoding} encoding")
            continue
        except Exception as e:
            logger.error(f"Error reading file with {encoding}: {str(e)}")
    
    # Last resort: read as binary and convert to string with replacement
    try:
        with open(file_path, 'rb') as file:
            binary_content = file.read()
            if binary_content.startswith(b'\xef\xbb\xbf'):  # BOM
                binary_content = binary_content[3:]
            
            # Try to decode with error replacement
            try:
                content = binary_content.decode('utf-8', errors='replace')
                logger.info("Successfully read file in binary mode with utf-8 and error replacement")
                return content
            except:
                # Final fallback
                logger.warning("Using last resort fallback for file content")
                return "Unable to extract text from this file. Please provide a plain text version."
    except Exception as e:
        logger.error(f"Error in binary file reading: {str(e)}")
        return "Error reading file. Please try a different format."

def generate_pdf(content, display_text):
    """Generate a PDF document from HTML content using ReportLab"""
    try:
        logger.info("Generating PDF with ReportLab")
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=20*mm,
            leftMargin=20*mm,
            topMargin=20*mm,
            bottomMargin=20*mm
        )
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'Title',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=10,
            textColor=colors.black
        )
        normal_style = ParagraphStyle(
            'Normal',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=5
        )
        
        elements = []
        
        # Clean HTML content for PDF rendering
        clean_content = content.replace('<h2>', '\n<h2>')
        clean_content = re.sub(r'<li>(.*?)</li>', r'• \1', clean_content)
        clean_content = re.sub(r'</?ul>', '', clean_content)
        clean_content = re.sub(r'<strong>(.*?)</strong>', r'\1', clean_content)
        
        for line in clean_content.split('\n'):
            line = line.strip()
            if not line:
                elements.append(Spacer(1, 5*mm))
                continue
            if '<h2>' in line:
                heading_text = re.sub(r'</?h2>', '', line).strip()
                elements.append(Paragraph(heading_text, title_style))
            else:
                clean_line = re.sub(r'<[^>]+>', '', line).strip()
                if clean_line:
                    elements.append(Paragraph(clean_line, normal_style))
        
        doc.build(elements)
        pdf_content = buffer.getvalue()
        buffer.close()
        
        logger.info(f"Successfully generated PDF of size {len(pdf_content)} bytes")
        return pdf_content, display_text
        
    except Exception as e:
        logger.error(f"PDF generation error: {str(e)}")
        
        # Create a simple PDF with error message as fallback
        try:
            logger.info("Attempting to create fallback PDF")
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            styles = getSampleStyleSheet()
            elements = [
                Paragraph("Resume Generation Error", styles['Heading1']),
                Paragraph("There was an error generating your PDF. Please try again.", styles['Normal']),
                Paragraph(f"Error details: {str(e)}", styles['Normal'])
            ]
            doc.build(elements)
            pdf_content = buffer.getvalue()
            buffer.close()
            logger.info("Successfully created fallback PDF")
            return pdf_content, "Error generating formatted PDF. Created a simple version instead."
        except Exception as fallback_error:
            logger.error(f"Fallback PDF generation error: {str(fallback_error)}")
            raise

def extract_text_from_docx(docx_path):
    """Extract text from a DOCX file"""
    try:
        logger.info(f"Extracting text from DOCX: {docx_path}")
        doc = docx.Document(docx_path)
        
        # Extract text from paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        
        content = "\n".join(paragraphs)
        logger.info(f"Successfully extracted {len(content)} characters from DOCX")
        return content
    except Exception as e:
        logger.error(f"Error processing DOCX: {str(e)}")
        logger.error(traceback.format_exc())
        raise ValueError(f"Error processing DOCX file: {str(e)}")

def read_file_content(file_path):
    """Read file content based on its extension"""
    file_ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_ext == '.docx':
            return extract_text_from_docx(file_path)
        elif file_ext == '.pdf':
            return extract_text_from_pdf(file_path)
        elif file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()
                logger.info(f"Successfully read TXT file: {len(content)} characters")
                return content
        else:
            error_message = f"Unsupported file extension: {file_ext}"
            logger.error(error_message)
            raise ValueError(error_message)
    except Exception as e:
        if not isinstance(e, ValueError):
            logger.error(f"Error reading file {file_path}: {str(e)}")
            logger.error(traceback.format_exc())
            raise ValueError(f"Error reading file: {str(e)}")
        raise

# Simplify and reduce complexity
def read_docx_or_txt(file_path):
    """Simplified function to read the content of a DOCX or TXT file"""
    ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if ext == '.docx':
            # Direct extraction of text from DOCX
            doc = docx.Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            content = "\n".join(paragraphs)
            logger.info(f"Read DOCX file successfully: {len(content)} characters")
            return content
        elif ext == '.txt':
            # Simple text file reading
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()
            logger.info(f"Read TXT file successfully: {len(content)} characters")
            return content
        else:
            raise ValueError(f"Unsupported file format: {ext}. Only .txt and .docx files are supported.")
    except Exception as e:
        logger.error(f"Error reading file {file_path}: {str(e)}")
        logger.error(traceback.format_exc())
        raise ValueError(f"Failed to read file: {str(e)}")

def generate_resume_with_extracted_text(resume_text, job_desc_text, template_text=None):
    max_retries = 3
    retry_delay = 2
    
    for attempt in range(max_retries):
        try:
            logger.info("Initializing Groq client")
            client = Groq()
            
            logger.info("Setting up messages for AI")
            messages = [
                {
                    "role": "system",
                    "content": """You are a professional resume writer assistant. 
Format your response with HTML tags for styling. Use <strong> for bold, <h2> for section headers, <ul> and <li> for lists. 
Provide ONLY the resume content, no analysis or insights. Start directly with the resume content.
When a template is provided, match its structure exactly, including headings, bullet points, and overall organization."""
                }
            ]

            if template_text:
                logger.info("Using template for formatting")
                
                prompt = f"""Create a tailored resume based on the provided resume content and job description, 
following EXACTLY the format and structure of the template document.

Match the template's exact organization, heading style, bullet points, and overall layout.
The goal is to make the generated resume look identical in structure to the template while updating the content
based on the provided resume and optimizing it for the job description.

Format the response with HTML tags for styling:
- Use <h2> for section headers
- Use <strong> for bold text
- Use <ul> and <li> for lists
- Maintain the same section order as the template

Resume Content:
{resume_text}

Job Description:
{job_desc_text}

Template to Match Exactly:
{template_text}

Remember: Match the exact structure, organization, and format of the template while tailoring the content to the job description."""

                messages.append({
                    "role": "user",
                    "content": prompt
                })
            else:
                messages.append({
                    "role": "user",
                    "content": f"Create a tailored resume based on this resume and job description. Format the response with HTML tags. Provide ONLY the resume content, no analysis or insights.\n\nResume:\n\n{resume_text}\n\nJob Description:\n\n{job_desc_text} and the template is {template_text}"
                })

            logger.info("Sending request to Groq API")
            full_response = ""
            try:
                logger.info("Trying non-streaming request")
                full_response_obj = client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=messages,
                    temperature=0.1,
                    max_tokens=2048,
                    top_p=1,
                    stream=False
                )
                full_response = full_response_obj.choices[0].message.content
                logger.info("Successfully received non-streaming response")
            except Exception as non_stream_error:
                logger.warning(f"Non-streaming request failed: {str(non_stream_error)}, trying streaming")
                completion = client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=messages,
                    temperature=0.1,
                    max_tokens=2048,
                    top_p=1,
                    stream=True
                )

                full_response = ""
                for chunk in completion:
                    if hasattr(chunk.choices[0], 'delta') and chunk.choices[0].delta.content is not None:
                        full_response += chunk.choices[0].delta.content
                logger.info("Successfully received streaming response")
            
            if not full_response:
                logger.error("Empty response received from Groq API")
                raise Exception("Empty response received from AI")
            
            logger.info(f"Received response of length {len(full_response)}")
            
            # Generate unique ID for this resume
            resume_id = str(int(time.time()))
            
            # Create display text
            template_info_text = ""
            if template_text:
                template_info_text = " following your provided template format"
                
            display_text = f"<p>Here's your tailored resume{template_info_text}:</p>\n\n{full_response}"
            
            # Generate PDF
            logger.info("Generating PDF from AI response")
            pdf_content, _ = generate_pdf(full_response, display_text)
            
            # Return both the PDF and the display text with resume ID
            return {
                'resume_id': resume_id,
                'pdf_content': pdf_content,
                'display_text': display_text,
                'timestamp': datetime.datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Error in generate_resume (attempt {attempt + 1}/{max_retries}): {str(e)}")
            if attempt < max_retries - 1:
                logger.info(f"Retrying in {retry_delay} seconds...")
                time.sleep(retry_delay)
                continue
            raise Exception(f"Failed to generate resume after {max_retries} attempts: {str(e)}")

class ResumeHandler(BaseHTTPRequestHandler):
    def _set_headers(self, content_type='text/plain'):
        self.send_response(200)
        self.send_header('Content-type', content_type)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.end_headers()

    def _handle_error(self, status_code, message):
        self.send_response(status_code)
        self.send_header('Content-type', 'application/json')
        self.send_header('Access-Control-Allow-Origin', '*')
        self.end_headers()
        error_response = json.dumps({'error': message})
        self.wfile.write(error_response.encode('utf-8'))

    def do_GET(self):
        parts = self.path.split('?')
        path = parts[0]
        
        if path == '/':
            try:
                self._set_headers('text/html')
                with open('index.html', 'rb') as file:
                    self.wfile.write(file.read())
            except Exception as e:
                self._handle_error(500, f"Error serving HTML: {str(e)}")
        elif path == '/history':
            try:
                self._set_headers('application/json')
                history = self._load_history()
                self.wfile.write(json.dumps(history).encode('utf-8'))
            except Exception as e:
                self._handle_error(500, f"Error loading history: {str(e)}")
        elif path.startswith('/download/'):
            try:
                # Extract resume ID from path
                resume_id = path.split('/download/')[1]
                logger.info(f"Download request for resume ID: {resume_id}")
                
                if not resume_id:
                    self._handle_error(400, "Missing resume ID")
                    return
                    
                history = self._load_history()
                if not history:
                    logger.error(f"History is empty")
                    self._handle_error(500, "No resumes in history")
                    return
                    
                logger.info(f"Looking for ID {resume_id} in history with {len(history)} items")
                resume_item = None
                for item in history:
                    if item.get('id') == resume_id:
                        resume_item = item
                        break
                        
                if not resume_item:
                    logger.error(f"Resume with ID {resume_id} not found in history")
                    self._handle_error(404, "Resume not found")
                    return
                    
                if 'resume' not in resume_item:
                    logger.error(f"Resume data missing for ID {resume_id}")
                    self._handle_error(404, "Resume data not found")
                    return
                
                # Force download with appropriate headers
                self.send_response(200)
                self.send_header('Content-Type', 'application/pdf')
                self.send_header('Content-Disposition', f'attachment; filename="resume_{resume_id}.pdf"')
                # Add critical headers for IE compatibility
                self.send_header('Cache-Control', 'no-cache, no-store, must-revalidate')
                self.send_header('Pragma', 'no-cache')
                self.send_header('Expires', '0')
                self.end_headers()
                
                # Decode base64 and send PDF content directly
                try:
                    pdf_data = resume_item.get('resume', '')
                    logger.info(f"Got base64 data of length: {len(pdf_data)}")
                    
                    # Ensure we are properly decoding the base64 content
                    pdf_content = base64.b64decode(pdf_data)
                    logger.info(f"Decoded PDF content of size: {len(pdf_content)} bytes")
                    
                    # Write the PDF bytes directly to the response
                    try:
                        self.wfile.write(pdf_content)
                        self.wfile.flush()
                        logger.info(f"Successfully sent PDF content")
                    except BrokenPipeError:
                        # This can happen if the client disconnects
                        logger.warning("Client disconnected during download")
                    except Exception as write_err:
                        logger.error(f"Error writing PDF data: {str(write_err)}")
                        raise
                        
                except Exception as pdf_error:
                    logger.error(f"PDF decode/write error: {str(pdf_error)}")
                    logger.error(traceback.format_exc())
                    self._handle_error(500, f"Error processing PDF: {str(pdf_error)}")
                
            except Exception as e:
                logger.error(f"Error serving PDF: {str(e)}")
                logger.error(traceback.format_exc())
                self._handle_error(500, f"Error serving PDF: {str(e)}")
        else:
            self._handle_error(404, "Not Found")

    def _load_history(self):
        try:
            if os.path.exists('resume_history.json'):
                logger.info("Loading history from file")
                with open('resume_history.json', 'r') as f:
                    history = json.load(f)
                    logger.info(f"Loaded {len(history)} history items")
                    return history
            logger.info("No history file found, returning empty list")
            return []
        except Exception as e:
            logger.error(f"Error loading history: {str(e)}")
            return []

    def _save_history(self, history):
        try:
            logger.info(f"Saving history with {len(history)} items")
            with open('resume_history.json', 'w') as f:
                json.dump(history, f, ensure_ascii=False)
        except Exception as e:
            logger.error(f"Error saving history: {str(e)}")

    def do_POST(self):
        if self.path == '/generate':
            try:
                temp_dir = tempfile.mkdtemp()
                logger.info(f"Created temporary directory: {temp_dir}")
                
                content_type = self.headers.get('Content-Type')
                if not content_type or not content_type.startswith('multipart/form-data'):
                    raise ValueError(f"Invalid Content-Type: {content_type}")
                
                # Get content length to read exact amount of data
                content_length = int(self.headers.get('Content-Length', 0))
                if content_length == 0:
                    raise ValueError("No content received")
                    
                # Read the exact amount of data
                post_data = self.rfile.read(content_length)
                
                # Create a cgi.FieldStorage object
                environ = {
                    'REQUEST_METHOD': 'POST',
                    'CONTENT_TYPE': content_type,
                    'CONTENT_LENGTH': str(content_length),
                }
                
                form = cgi.FieldStorage(
                    fp=io.BytesIO(post_data),
                    headers=self.headers,
                    environ=environ
                )

                if 'resume' not in form or 'job_desc' not in form:
                    raise ValueError("Missing required files: Both 'resume' and 'job_desc' are required")

                resume_file = form['resume']
                job_desc_file = form['job_desc']

                if not resume_file.filename or not job_desc_file.filename:
                    raise ValueError("Empty file names provided")

                resume_ext = os.path.splitext(resume_file.filename)[1].lower()
                job_desc_ext = os.path.splitext(job_desc_file.filename)[1].lower()
                
                # Validate file types
                allowed_exts = ['.txt', '.docx']
                if resume_ext not in allowed_exts:
                    raise ValueError(f"Resume must be a text (.txt) or Word (.docx) file, got {resume_ext}")
                    
                if job_desc_ext not in allowed_exts:
                    raise ValueError(f"Job description must be a text (.txt) or Word (.docx) file, got {job_desc_ext}")
                
                resume_path = os.path.join(temp_dir, f'resume{resume_ext}')
                job_desc_path = os.path.join(temp_dir, f'job_desc{job_desc_ext}')

                # Save files in binary mode
                with open(resume_path, 'wb') as f:
                    if hasattr(resume_file, 'file'):
                        f.write(resume_file.file.read())
                    else:
                        f.write(resume_file.value)
                    logger.info(f"Saved resume file to {resume_path}")
                    
                with open(job_desc_path, 'wb') as f:
                    if hasattr(job_desc_file, 'file'):
                        f.write(job_desc_file.file.read())
                    else:
                        f.write(job_desc_file.value)
                    logger.info(f"Saved job description file to {job_desc_path}")

                template_path = None
                template_text = None
                if 'template' in form and form['template'].filename:
                    template_file = form['template']
                    template_ext = os.path.splitext(template_file.filename)[1].lower()
                    
                    # Only allow text and docx files for template
                    if template_ext not in allowed_exts:
                        raise ValueError(f"Template must be a text (.txt) or Word (.docx) file, got {template_ext}")
                        
                    template_path = os.path.join(temp_dir, f'template{template_ext}')
                    with open(template_path, 'wb') as f:
                        if hasattr(template_file, 'file'):
                            f.write(template_file.file.read())
                        else:
                            f.write(template_file.value)
                        logger.info(f"Saved template file to {template_path}")
                    
                    # Use simplified file reading
                    try:
                        logger.info(f"Reading template file: {template_path}")
                        template_text = read_docx_or_txt(template_path)
                        if not template_text or template_text.strip() == "":
                            raise ValueError("Template file is empty")
                        logger.info(f"Template text length: {len(template_text)}")
                    except Exception as e:
                        logger.error(f"Failed to read template file: {str(e)}")
                        raise ValueError(f"Failed to read template: {str(e)}")

                # Read the files
                logger.info("Reading resume file")
                resume_text = read_docx_or_txt(resume_path)
                
                logger.info("Reading job description file")
                job_desc_text = read_docx_or_txt(job_desc_path)
                
                logger.info("Generating resume")
                result = generate_resume_with_extracted_text(
                    resume_text, job_desc_text, template_text
                )

                # Save to history with the PDF content
                history = self._load_history()
                history.append({
                    'id': result['resume_id'],
                    'timestamp': result['timestamp'],
                    'display_text': result['display_text'],
                    'resume': base64.b64encode(result['pdf_content']).decode('utf-8')
                })
                self._save_history(history)
                logger.info(f"Saved resume {result['resume_id']} to history")

                response = {
                    'resume_id': result['resume_id'],
                    'resume': base64.b64encode(result['pdf_content']).decode('utf-8'),
                    'display_text': result['display_text'],
                    'mime_type': 'application/pdf'
                }
                
                self._set_headers('application/json')
                self.wfile.write(json.dumps(response, ensure_ascii=False).encode('utf-8'))
                logger.info(f"Successfully sent response for resume {result['resume_id']}")

            except ValueError as e:
                logger.error(f"Value error: {str(e)}")
                self._handle_error(400, str(e))
            except Exception as e:
                logger.error(f"Error in do_POST: {str(e)}")
                logger.error(traceback.format_exc())
                self._handle_error(500, str(e))
            finally:
                if 'temp_dir' in locals() and os.path.exists(temp_dir):
                    try:
                        shutil.rmtree(temp_dir)
                        logger.info(f"Removed temporary directory: {temp_dir}")
                    except Exception as e:
                        logger.error(f"Error removing temp dir: {str(e)}")
        elif self.path == '/delete':
            try:
                content_length = int(self.headers.get('Content-Length', 0))
                data = json.loads(self.rfile.read(content_length))
                resume_id = data.get('resume_id')
                
                if not resume_id:
                    raise ValueError("Missing resume_id")
                
                history = self._load_history()
                history = [item for item in history if item['id'] != resume_id]
                self._save_history(history)
                
                self._set_headers('application/json')
                self.wfile.write(json.dumps({'success': True}).encode('utf-8'))
            except Exception as e:
                self._handle_error(500, str(e))
        else:
            self._handle_error(404, "Not Found")

if __name__ == '__main__':
    host = 'localhost'
    port = 8000
    server = HTTPServer((host, port), ResumeHandler)
    logger.info(f'Server running on http://{host}:{port}')
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        logger.info('Server stopping...')
    except Exception as e:
        logger.error(f'Server error: {str(e)}')
    finally:
        server.server_close()
        logger.info('Server stopped')