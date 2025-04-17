"""
Test script to verify DOCX file reading and handling
"""
import os
import docx
import io
import base64

def create_test_docx(filename="test_template.docx"):
    """Create a test DOCX file with typical resume sections"""
    doc = docx.Document()
    
    # Add a title
    doc.add_heading('Test Resume Template', 0)
    
    # Personal Info section
    doc.add_heading('Personal Information', 1)
    doc.add_paragraph('Name: John Doe')
    doc.add_paragraph('Email: john.doe@example.com')
    doc.add_paragraph('Phone: (123) 456-7890')
    
    # Summary section
    doc.add_heading('Professional Summary', 1)
    doc.add_paragraph('Experienced professional with skills in various technologies...')
    
    # Skills section
    doc.add_heading('Skills', 1)
    skills = doc.add_paragraph()
    skills.add_run('Technical Skills: ').bold = True
    skills.add_run('Python, JavaScript, HTML, CSS')
    
    doc.add_paragraph().add_run().add_break()
    
    skills2 = doc.add_paragraph()
    skills2.add_run('Soft Skills: ').bold = True
    skills2.add_run('Communication, Leadership, Problem Solving')
    
    # Experience section
    doc.add_heading('Work Experience', 1)
    
    # Job 1
    job1 = doc.add_paragraph()
    job1.add_run('Software Developer').bold = True
    job1.add_run(' | ABC Company | 2018-2022')
    
    # Add bullet points
    doc.add_paragraph('• Developed web applications using modern technologies', style='List Bullet')
    doc.add_paragraph('• Collaborated with cross-functional teams', style='List Bullet')
    doc.add_paragraph('• Implemented new features and fixed bugs', style='List Bullet')
    
    # Job 2
    job2 = doc.add_paragraph()
    job2.add_run('Junior Developer').bold = True
    job2.add_run(' | XYZ Corp | 2016-2018')
    
    doc.add_paragraph('• Assisted senior developers with coding tasks', style='List Bullet')
    doc.add_paragraph('• Participated in code reviews and testing', style='List Bullet')
    
    # Education section
    doc.add_heading('Education', 1)
    edu = doc.add_paragraph()
    edu.add_run('Bachelor of Science in Computer Science').bold = True
    edu.add_run(' | University Name | 2016')
    
    # Save the document
    doc.save(filename)
    print(f"Created test DOCX file: {filename}")
    return filename

def read_docx_file(filename):
    """Read and analyze a DOCX file"""
    try:
        doc = docx.Document(filename)
        print(f"\nSuccessfully opened DOCX file: {filename}")
        
        # Basic stats
        print(f"Document has {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
        
        # Print headings
        headings = []
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                headings.append(para.text)
        
        print("\nHeadings:")
        for heading in headings:
            print(f"  - {heading}")
        
        # Print some content
        print("\nSample paragraphs:")
        for i, para in enumerate(doc.paragraphs[:5]):
            if para.text.strip():
                print(f"  {i+1}. {para.text[:50]}...")
                
        return True
    except Exception as e:
        print(f"Error reading DOCX file: {str(e)}")
        return False

def main():
    """Main function to test DOCX functionality"""
    print("Testing python-docx functionality")
    print("=================================")
    
    # Create a test DOCX file
    test_file = create_test_docx()
    
    # Read the DOCX file
    success = read_docx_file(test_file)
    
    if success:
        print("\nTest completed successfully!")
    else:
        print("\nTest failed!")

if __name__ == "__main__":
    main() 